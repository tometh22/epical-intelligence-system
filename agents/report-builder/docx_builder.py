"""Word document (.docx) generation for intelligence reports."""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from agents.shared.logger import get_logger

logger = get_logger("report-builder")


def build_report_docx(
    client_name: str,
    period: str,
    report_text: str,
    metrics: Dict[str, Any],
    output_path: Union[str, Path],
) -> Path:
    """Create a professionally formatted .docx intelligence report.

    Args:
        client_name: Client name for the title page.
        period: Reporting period description.
        report_text: The Claude-generated narrative text.
        metrics: Calculated metrics dictionary for the summary table.
        output_path: Destination file path.

    Returns:
        Path to the generated .docx file.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Title page --
    _add_title_page(doc, client_name, period)

    # -- Table of contents placeholder --
    doc.add_page_break()
    toc_heading = doc.add_heading("Tabla de Contenidos", level=1)
    toc_para = doc.add_paragraph(
        "[Actualizar este campo en Word: References > Update Table of Contents]"
    )
    toc_para.italic = True
    doc.add_page_break()

    # -- Metrics summary table --
    doc.add_heading("Resumen de Metricas", level=1)
    _add_metrics_table(doc, metrics)
    doc.add_paragraph("")  # spacer

    # -- Actor breakdown table --
    actor_breakdown = metrics.get("actor_breakdown")
    if actor_breakdown:
        doc.add_heading("Desglose por Actor", level=2)
        actor_total = sum(actor_breakdown.values()) or 1
        actor_table = doc.add_table(
            rows=len(actor_breakdown) + 1, cols=3, style="Light Shading Accent 1",
        )
        actor_table.rows[0].cells[0].text = "Actor"
        actor_table.rows[0].cells[1].text = "Menciones"
        actor_table.rows[0].cells[2].text = "Porcentaje"
        for i, (actor_name, actor_count) in enumerate(actor_breakdown.items(), 1):
            actor_table.rows[i].cells[0].text = str(actor_name)
            actor_table.rows[i].cells[1].text = f"{actor_count:,}"
            actor_table.rows[i].cells[2].text = f"{round(actor_count / actor_total * 100, 1)}%"
        doc.add_paragraph("")  # spacer

    # -- Intersection summary --
    inter = metrics.get("intersection")
    if inter and inter.get("total", 0) > 0:
        doc.add_heading("Análisis de Intersección", level=2)
        doc.add_paragraph(
            f"Análisis de intersección: {inter.get('brand_only', 0)} menciones exclusivas "
            f"de la marca ({inter.get('brand_only_pct', 0)}%), "
            f"{inter.get('actor_only', 0)} del actor secundario "
            f"({inter.get('actor_only_pct', 0)}%), "
            f"{inter.get('intersection', 0)} menciones que referencian ambos "
            f"({inter.get('intersection_pct', 0)}%)."
        )
        doc.add_paragraph("")  # spacer

    # -- Topic clusters --
    topic_clusters = metrics.get("topic_clusters")
    if topic_clusters:
        doc.add_heading("Clusters Temáticos", level=2)
        for cluster in topic_clusters:
            label = cluster.get("label", "Sin etiqueta")
            pct = cluster.get("percentage", 0)
            doc.add_paragraph(
                f"{label}: {pct}%", style="List Bullet",
            )
        doc.add_paragraph("")  # spacer

    # -- Anomalies --
    anomalies_list = metrics.get("anomalies_list")
    if anomalies_list:
        doc.add_heading("Anomalías Detectadas", level=2)
        for anomaly in anomalies_list:
            severity = anomaly.get("severity", "unknown").upper()
            anom_type = anomaly.get("type", "")
            description = anomaly.get("description", "")
            para = doc.add_paragraph()
            run_sev = para.add_run(f"[{severity}] {anom_type}: ")
            run_sev.bold = True
            para.add_run(description)
        doc.add_paragraph("")  # spacer

    # -- Report narrative sections --
    _add_report_sections(doc, report_text)

    # -- Footer --
    _add_footer(doc)

    doc.save(str(output_path))
    logger.info("DOCX report saved to %s", output_path)
    return output_path


def _add_title_page(doc: Document, client_name: str, period: str) -> None:
    """Add a title page with client name, subtitle, and generation date."""
    # Add some spacing
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(client_name)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informe de Inteligencia Social")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

    period_para = doc.add_paragraph()
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = period_para.add_run(period)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x8A)

    doc.add_paragraph("")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generado: {datetime.now().strftime('%d de %B de %Y')}")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_metrics_table(doc: Document, metrics: Dict[str, Any]) -> None:
    """Add a formatted table with key metrics."""
    rows_data: List[Tuple[str, str]] = [
        ("Total de menciones", str(metrics.get("total_mentions", "N/D"))),
    ]

    # Sentiment breakdown
    sentiment = metrics.get("sentiment_breakdown", {})
    if sentiment:
        parts = []
        for key, val in sentiment.items():
            if isinstance(val, dict):
                parts.append(f"{key}: {val.get('count', 0)} ({val.get('percentage', 0)}%)")
            else:
                parts.append(f"{key}: {val}")
        rows_data.append(("Sentimiento", " | ".join(parts)))

    # Top sources
    top_sources = metrics.get("top_sources", [])
    if top_sources:
        sources_str = ", ".join(f"{s} ({c})" for s, c in top_sources[:5])
        rows_data.append(("Principales fuentes", sources_str))

    if metrics.get("avg_reach") is not None:
        rows_data.append(("Alcance promedio", f"{metrics['avg_reach']:,.0f}"))

    if metrics.get("avg_engagement") is not None:
        rows_data.append(("Engagement promedio", f"{metrics['avg_engagement']:,.1f}"))

    table = doc.add_table(rows=len(rows_data), cols=2, style="Light Shading Accent 1")
    for i, (label, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value


def _add_report_sections(doc: Document, report_text: str) -> None:
    """Parse the report narrative and add sections with proper heading styles."""
    lines = report_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Detect markdown-style headings
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
            # Numbered list items
            text = stripped.split(". ", 1)[1] if ". " in stripped else stripped
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(stripped)


def _add_footer(doc: Document) -> None:
    """Add a footer to all sections of the document."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            "Generado por Epical Intelligence System \u2014 Borrador para revision"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
